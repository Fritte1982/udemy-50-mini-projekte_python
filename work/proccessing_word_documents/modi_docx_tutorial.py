import docx

from own_packages import path_attributes
from pathlib import Path
from docx.shared import RGBColor

output_path = path_attributes.OUTPUT_PATH
word_sources_path = path_attributes.DOCX_SOURCE
docpath1 = word_sources_path / "panda1.docx"
docpath2 = word_sources_path / "panda2.docx"
docpath_fin = Path(output_path) / "panda_fin.docx"

doc1 = docx.Document(docpath1)
doc2 = docx.Document(docpath2)


para_doc2 = doc2.paragraphs[0]

# paragraph_obj_doc1 = doc1.add_paragraph(text_doc2)
paragraphs = doc1.paragraphs
paragraphs[1]._element.addnext(para_doc2._element)
doc1.save(str(docpath_fin))

doc_fin = docx.Document(docpath_fin)
para_doc_fin = doc_fin.paragraphs[2]

# Textinhalt sichern
text = para_doc_fin.text

# Paragraph löschen: geht in python-docx leider nicht direkt,
# aber du kannst ihn einfach "leeren", indem du ihn ersetzt:
p = para_doc_fin._element
for child in list(p):
    p.remove(child)

# Neuen Run hinzufügen mit neu gesetzter Farbe
new_run = para_doc_fin.add_run(text)
new_run.font.color.rgb = RGBColor(0, 0, 0)  # Jetzt wirklich schwarz!

# Dokument speichern
doc_fin.save(str(docpath_fin))




