from pathlib import Path
from own_packages import path_attributes as pat
import docx

words_path = Path(pat.DOCX_SOURCE)
output_doc = Path(pat.OUTPUT_PATH) / "merged_a.docx"

merged_document = docx.Document()
for path in words_path.iterdir():
    if str(path.name).startswith("sam"):
        doc = docx.Document(str(path))
        for paragraph in doc.paragraphs:
            new_para =merged_document.add_paragraph()
            new_para._element.addprevious(paragraph._element)
merged_document.save(str(output_doc))
