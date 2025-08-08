import docx
from docx.shared import RGBColor
from own_packages import path_attributes
from pprint import pprint

DOCX_IN_1 = path_attributes.DOCX_SOURCE / "panda1.docx"
DOCX_IN_2 = path_attributes.DOCX_SOURCE / "panda2.docx"

out_doc = path_attributes.OUTPUT_PATH / "panda_merge.docx"

if DOCX_IN_1.exists() and DOCX_IN_2.exists():
    print("Panda-docx's exists")

doc_pandas1: "docx.Document" = docx.Document(DOCX_IN_1)
doc_pandas2: "docx.Document" = docx.Document(DOCX_IN_2)

black_color = RGBColor(0,0,0)

# para_source = doc_pandas2.paragraphs[0]
# para_target = doc_pandas1.add_paragraph()
# for run in para_source.runs:
#     new_run = para_target.add_run(run.text)
#     new_run.font.color.rgb = black_color
# # para.style.font.color.rgb = black_color
# # new_pargraph =  doc_pandas1.add_paragraph(txt)
# paragraphs = doc_pandas1.paragraphs
# paragraphs[1]._element.addnext(para_target._element)
# doc_pandas1.save(out_doc)

black_color = RGBColor(0, 0, 0)

# Nur den Text übernehmen
txt = doc_pandas2.paragraphs[0].text

# Absatz im doc_pandas1 erzeugen → übernimmt Standardstil von doc_pandas1
new_para = doc_pandas1.add_paragraph(txt)

# Farbe direkt im neuen Absatz setzen
for run in new_para.runs:
    run.font.color.rgb = black_color

# An gewünschte Position verschieben
doc_pandas1.paragraphs[1]._element.addnext(new_para._element)

doc_pandas1.save(out_doc)



for para in doc_pandas1.paragraphs:
    print(type(para.text))
    pprint(para.text)

